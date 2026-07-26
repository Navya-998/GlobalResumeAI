from docx import Document
from docx.shared import Pt


# ==========================================================
# Resume DOCX
# ==========================================================

def generate_docx(resume):

    doc = Document()

    title = doc.add_heading(
        resume.full_name,
        level=1
    )

    title.style.font.size = Pt(20)

    contact = []

    if resume.email:
        contact.append(resume.email)

    if resume.phone:
        contact.append(resume.phone)

    if resume.address:
        contact.append(resume.address)

    if resume.linkedin:
        contact.append(f"LinkedIn: {resume.linkedin}")

    if resume.github:
        contact.append(f"GitHub: {resume.github}")

    doc.add_paragraph(" | ".join(contact))

    if resume.summary:
        doc.add_heading(
            "Professional Summary",
            level=2
        )
        doc.add_paragraph(
            resume.summary
        )

    if resume.education:
        doc.add_heading(
            "Education",
            level=2
        )
        doc.add_paragraph(
            resume.education
        )

    if resume.experience:
        doc.add_heading(
            "Experience",
            level=2
        )
        doc.add_paragraph(
            resume.experience
        )

    if resume.skills:
        doc.add_heading(
            "Skills",
            level=2
        )
        doc.add_paragraph(
            resume.skills
        )

    if resume.projects:
        doc.add_heading(
            "Projects",
            level=2
        )
        doc.add_paragraph(
            resume.projects
        )

    if resume.certifications:
        doc.add_heading(
            "Certifications",
            level=2
        )
        doc.add_paragraph(
            resume.certifications
        )

    if resume.custom_sections:
        doc.add_heading(
            "Additional Information",
            level=2
        )
        doc.add_paragraph(
            resume.custom_sections
        )

    return doc


# ==========================================================
# Cover Letter DOCX
# ==========================================================

def generate_cover_letter(resume):

    doc = Document()

    heading = doc.add_heading(
        "Cover Letter",
        level=1
    )

    heading.style.font.size = Pt(20)

    doc.add_paragraph()

    doc.add_paragraph("Dear Hiring Manager,")

    doc.add_paragraph()

    if resume.summary:

        doc.add_paragraph(
            resume.summary
        )

    else:

        doc.add_paragraph(

            f"My name is {resume.full_name}. "
            "I am excited to apply for this opportunity. "
            "My skills and experience make me a strong candidate "
            "for this position."

        )

    if resume.skills:

        doc.add_heading(
            "Key Skills",
            level=2
        )

        doc.add_paragraph(
            resume.skills
        )

    if resume.experience:

        doc.add_heading(
            "Professional Experience",
            level=2
        )

        doc.add_paragraph(
            resume.experience
        )

    doc.add_paragraph()

    doc.add_paragraph(

        "I would welcome the opportunity to discuss "
        "how my knowledge and experience can contribute "
        "to your organization."

    )

    doc.add_paragraph()

    doc.add_paragraph(

        "Thank you for your time and consideration."

    )

    doc.add_paragraph()

    doc.add_paragraph("Sincerely,")

    doc.add_paragraph(
        resume.full_name
    )

    return doc